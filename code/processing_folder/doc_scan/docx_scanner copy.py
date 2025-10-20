import os, sys, json, hashlib, pickle, re
from typing import List, Tuple, Iterable, Dict, Optional, Any
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

# your helpers
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from excel_scan.embedding_helper import get_embedding, get_cosine_similarity


# ---------------------------- Config ---------------------------------

DEFAULT_THRESHOLD = 0.35
DEFAULT_TOP_K     = 10
MAX_CHARS_PER_CHUNK = 500
EMB_CACHE_DIR = Path(__file__).parent / ".emb_cache"

# Owner header families we support
OWNER_PREFIXES = ["CNTR", "API", "EDI", "VSL", "VESSEL", "VSSL"]
OWNER_PREFIX_RE = re.compile(r"^\s*(" + "|".join(OWNER_PREFIXES) + r")\s*:\s*", re.I)

SUBSECTION_TITLES = ["Overview", "Resolution", "Verification"]
SUBSECTION_TITLE_RE = re.compile(r"^\s*(overview|resolution|verification)\s*:?\s*$", re.I)

# ---------------------------------------------------------------------


def _normalize_text(s: str) -> str:
    s = (s or "").replace("\r", " ").strip()
    while "  " in s:
        s = s.replace("  ", " ")
    return s


def _chunk_text(s: str, max_chars: int = MAX_CHARS_PER_CHUNK) -> List[str]:
    s = _normalize_text(s)
    if len(s) <= max_chars:
        return [s] if s else []
    chunks, start = [], 0
    while start < len(s):
        end = min(start + max_chars, len(s))
        cut = s.rfind(". ", start, end)
        if cut == -1 or cut < start + int(max_chars * 0.5):
            cut = end
        else:
            cut += 2
        chunks.append(s[start:cut].strip())
        start = cut
    return [c for c in chunks if c]


def _iter_docx_texts(docx_path: str, include_tables: bool = True,
                     chunk: bool = True, max_chars: int = MAX_CHARS_PER_CHUNK
                     ) -> Iterable[Tuple[str, str]]:
    """
    Yields (source_id, text) from the .docx file:
      - Paragraphs as P#<idx>
      - (Optional) Table cells as T<table_idx>R<row_idx>C<col_idx>
    """
    doc = Document(docx_path)

    # Paragraphs
    for i, p in enumerate(doc.paragraphs):
        t = _normalize_text(p.text or "")
        if not t:
            continue
        if chunk:
            for j, ch in enumerate(_chunk_text(t, max_chars=max_chars)):
                yield (f"P#{i}:{j}", ch)
        else:
            yield (f"P#{i}", t)

    # Tables
    if include_tables:
        for ti, tbl in enumerate(doc.tables):
            for ri, row in enumerate(tbl.rows):
                for ci, cell in enumerate(row.cells):
                    t = _normalize_text(cell.text or "")
                    if not t:
                        continue
                    if chunk:
                        for j, ch in enumerate(_chunk_text(t, max_chars=max_chars)):
                            yield (f"T{ti}R{ri}C{ci}:{j}", ch)
                    else:
                        yield (f"T{ti}R{ri}C{ci}", t)


# ------------------------ Embedding cache -----------------------------

def _sha1(text: str) -> str:
    return hashlib.sha1(text.encode("utf-8")).hexdigest()

def _load_cache() -> Dict[str, List[float]]:
    EMB_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    pkl = EMB_CACHE_DIR / "embeddings.pkl"
    if pkl.exists():
        try:
            with open(pkl, "rb") as f:
                return pickle.load(f)
        except Exception:
            return {}
    return {}

def _save_cache(cache: Dict[str, List[float]]) -> None:
    pkl = EMB_CACHE_DIR / "embeddings.pkl"
    with open(pkl, "wb") as f:
        pickle.dump(cache, f)

def _embed_with_cache(text: str, cache: Dict[str, List[float]]) -> Optional[List[float]]:
    key = _sha1(text)
    if key in cache:
        return cache[key]
    emb = get_embedding(text)
    if emb:
        cache[key] = emb
    return emb


# ------------------------ Main search API -----------------------------

def semantic_docx_search(
    query: str,
    docx_path: str,
    threshold: float = DEFAULT_THRESHOLD,
    top_k: int = DEFAULT_TOP_K,
    include_tables: bool = True,
    chunk: bool = True,
    max_chars: int = MAX_CHARS_PER_CHUNK,
    return_always_topk: bool = True,
    keyword_fallback: bool = True,
) -> List[Tuple[str, str, float]]:
    if not os.path.exists(docx_path):
        return []

    cache = _load_cache()

    query_norm = _normalize_text(query)
    query_emb = get_embedding(query_norm)
    if not query_emb:
        return []

    matches: List[Tuple[str, str, float]] = []
    seen_text: Dict[str, float] = {}

    for source_id, text in _iter_docx_texts(
        docx_path, include_tables=include_tables, chunk=chunk, max_chars=max_chars
    ):
        if not text:
            continue
        if text in seen_text:
            sim = seen_text[text]
        else:
            emb = _embed_with_cache(text, cache)
            if not emb:
                continue
            sim = get_cosine_similarity(query_emb, emb)
            seen_text[text] = sim
        matches.append((source_id, text, sim))

    _save_cache(cache)

    matches.sort(key=lambda x: x[2], reverse=True)
    filtered = [m for m in matches if m[2] >= threshold]
    if filtered:
        return filtered[:top_k] if top_k else filtered
    if return_always_topk and matches:
        return matches[:top_k] if top_k else matches

    if keyword_fallback:
        q_words = set(w.lower() for w in query_norm.split() if len(w) > 3)
        kmatches = []
        for sid, text, _sim in matches:
            tset = set(text.lower().split())
            overlap = len(q_words & tset)
            if overlap >= 2:
                kmatches.append((sid, text, 0.0))
        if kmatches:
            return kmatches[:top_k] if top_k else kmatches

    return []


# ------------------------ Structure helpers ---------------------------

def _heading_level(p: Paragraph) -> Optional[int]:
    """Returns heading level (1..9) if paragraph style is 'Heading N'; else None."""
    try:
        name = (p.style.name or "").strip()
    except Exception:
        name = ""
    m = re.match(r"Heading\s+([1-9])$", name, flags=re.I)
    return int(m.group(1)) if m else None

def _is_list_item(p: Paragraph) -> bool:
    """Heuristic: Word numbering (numPr) OR a visible bullet/number prefix."""
    try:
        numPr = p._p.pPr.numPr  # type: ignore[attr-defined]
        if numPr is not None:
            return True
    except Exception:
        pass
    t = (p.text or "").strip()
    return bool(re.match(r"^([\-–•●◦·]|(\d+|[A-Za-z])[\.\)]|•)\s+", t))

def iter_block_items(doc: Document):
    """Yields blocks in document order as ('p', Paragraph) or ('tbl', Table)."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith('p'):
            yield ('p', Paragraph(child, doc))
        elif child.tag.endswith('tbl'):
            yield ('tbl', Table(child, doc))

def build_doc_index(doc: Document) -> Dict[str, int]:
    """Map 'P#<i>' and 'T<i>' to block index."""
    idx: Dict[str, int] = {}
    p_counter, t_counter, b_counter = 0, 0, 0
    for kind, _obj in iter_block_items(doc):
        if kind == 'p':
            idx[f"P#{p_counter}"] = b_counter
            p_counter += 1
        else:
            idx[f"T{t_counter}"] = b_counter
            t_counter += 1
        b_counter += 1
    return idx

def _parse_source_id_root(source_id: str) -> Optional[str]:
    """Normalize source_id to paragraph/table root id used in build_doc_index."""
    m = re.match(r"^(P#\d+)(?::\d+)?$", source_id)
    if m: return m.group(1)
    m = re.match(r"^(T\d+)R\d+C\d+(?::\d+)?$", source_id)
    if m: return m.group(1)
    m = re.match(r"^(T\d+)(?::\d+)?$", source_id)
    if m: return m.group(1)
    return None


# ---------- New heuristics: owner/subsection recognition by text ------

def _is_owner_title(text: str) -> bool:
    return bool(OWNER_PREFIX_RE.match(text or ""))

def _owner_prefix(text: str) -> Optional[str]:
    m = OWNER_PREFIX_RE.match(text or "")
    return m.group(1).upper() if m else None

def _is_subsection_title(text: str) -> Optional[str]:
    m = SUBSECTION_TITLE_RE.match(text or "")
    return m.group(1).capitalize() if m else None


def _find_owner_bounds(blocks: List[Tuple[str, Any]], hit_block: int, max_scan: int = 1000) -> Tuple[int, int, Optional[int]]:
    """
    Find owner start as the nearest paragraph above that is either:
      - a styled Heading, OR
      - a line matching OWNER_PREFIX_RE (e.g., 'API: ...', 'CNTR: ...')
    Owner end is:
      - the next styled heading with level <= owner level (if owner is styled), OR
      - the next paragraph that matches OWNER_PREFIX_RE (if owner is text-only).
    """
    # 1) search upward
    start = None
    owner_level = None
    for i in range(hit_block, max(-1, hit_block - max_scan), -1):
        kind, obj = blocks[i]
        if kind != 'p':
            continue
        t = (obj.text or "").strip()
        lvl = _heading_level(obj)
        if lvl is not None:
            start, owner_level = i, lvl
            break
        if _is_owner_title(t):
            start, owner_level = i, None  # text-only owner
            break
    if start is None:
        start, owner_level = hit_block, None

    # 2) search forward for end
    end = len(blocks) - 1
    if owner_level is not None:
        # styled heading: end at next heading of <= level OR next owner title
        for k in range(start + 1, min(len(blocks), start + 1 + max_scan)):
            kind, obj = blocks[k]
            if kind != 'p':
                continue
            t = (obj.text or "").strip()
            lvl = _heading_level(obj)
            if lvl is not None and lvl <= owner_level:
                end = k - 1
                break
            if _is_owner_title(t):
                end = k - 1
                break
    else:
        # text-only owner: end at next owner title or next styled Heading 1/2
        for k in range(start + 1, min(len(blocks), start + 1 + max_scan)):
            kind, obj = blocks[k]
            if kind != 'p':
                continue
            t = (obj.text or "").strip()
            if _is_owner_title(t):
                end = k - 1
                break
            lvl = _heading_level(obj)
            if lvl is not None and lvl <= 2:  # conservative cut
                end = k - 1
                break

    return start, end, owner_level


def _extract_module_near_owner(blocks: List[Tuple[str, Any]], start: int, end: int) -> Optional[str]:
    """
    Look for a small table under the owner where first cell contains 'Module'
    and the adjacent/right cell has the desired value.
    """
    # scan first few blocks after owner for a table
    for i in range(start + 1, min(end + 1, start + 12)):
        kind, obj = blocks[i]
        if kind != 'tbl':
            continue
        tbl: Table = obj
        for r in tbl.rows:
            cells = [c.text.strip() for c in r.cells]
            if not cells:
                continue
            # common two-column "key | value" shape
            if len(cells) >= 2 and re.search(r"^module$", cells[0], re.I):
                val = cells[1].strip()
                return val or None
            # sometimes 'Module' might be elsewhere in the row
            for ci, cell in enumerate(cells):
                if re.search(r"^module$", cell, re.I) and ci + 1 < len(cells):
                    return cells[ci + 1].strip() or None
    return None


def _collect_between(blocks, start_idx, end_idx, *, lists_only=False, include_tables=True) -> List[str]:
    """Collect non-heading lines between indices (inclusive)."""
    lines: List[str] = []
    j = start_idx
    while j <= end_idx:
        kind, obj = blocks[j]
        if kind == 'p':
            # skip any heading lines
            if _heading_level(obj) is not None:
                j += 1
                continue
            t = (obj.text or "").strip()
            if not t:
                j += 1
                continue
            if (not lists_only) or _is_list_item(obj):
                lines.append(t)
        else:
            if include_tables:
                tbl: Table = obj
                for r in tbl.rows:
                    cells_text = [c.text.strip() for c in r.cells]
                    line = " | ".join([x for x in cells_text if x])
                    if line:
                        lines.append(line)
        j += 1
    return lines


def extract_owner_with_sections(
    docx_path: str,
    hit_source_id: str,
    wanted_sections: List[str] = SUBSECTION_TITLES,
    list_only_for: Dict[str, bool] = {"Overview": False, "Resolution": True, "Verification": True},
    include_tables_as_lines: bool = True,
    max_blocks_scan: int = 1000,
) -> Dict[str, Any]:
    """
    Robust extractor:
      - Owner header by style OR 'CNTR|API|EDI|VSL:'
      - Owner ends at next owner header or heading peer
      - Inside owner, find Overview/Resolution/Verification by TEXT match (not style)
    """
    doc = Document(docx_path)
    blocks: List[Tuple[str, Any]] = list(iter_block_items(doc))
    block_index = build_doc_index(doc)

    root_id = _parse_source_id_root(hit_source_id)
    if not root_id or root_id not in block_index:
        return {"owner_title": None, "owner_prefix": None, "owner_family": None, "owner_level": None, "module": None, "sections": [], "bounds": [None, None]}

    hit_block = block_index[root_id]
    owner_start, owner_end, owner_level = _find_owner_bounds(blocks, hit_block, max_scan=max_blocks_scan)

    # owner title & family
    owner_title = (blocks[owner_start][1].text or "").strip() if blocks[owner_start][0] == 'p' else ""
    owner_pref = _owner_prefix(owner_title)
    owner_family = None
    if owner_pref:
        up = owner_pref.upper()
        if up in {"VSL", "VSSL", "VESSEL"}: owner_family = "VSL"
        elif up in {"CNTR"}: owner_family = "CNTR"
        elif up in {"API"}: owner_family = "API"
        elif up in {"EDI"}: owner_family = "EDI"
        else: owner_family = up
    # module from nearby table
    module_val = _extract_module_near_owner(blocks, owner_start, owner_end)

    # find subheading starts by TEXT match
    sub_starts: Dict[str, int] = {}
    for idx in range(owner_start + 1, owner_end + 1):
        kind, obj = blocks[idx]
        if kind != 'p':
            continue
        title = _is_subsection_title((obj.text or "").strip())
        if title and title not in sub_starts:
            sub_starts[title] = idx

    # build sections in requested order
    found_sections: List[Dict[str, Any]] = []
    order = [w for w in wanted_sections if w in SUBSECTION_TITLES]
    for i, name in enumerate(order):
        if name not in sub_starts:
            continue
        start_idx = sub_starts[name] + 1
        # end = next subheading start - 1, else owner_end
        next_end = owner_end
        for j in range(start_idx, owner_end + 1):
            kind, obj = blocks[j]
            if kind != 'p':
                continue
            tname = _is_subsection_title((obj.text or "").strip())
            if tname and tname != name:
                next_end = j - 1
                break
        lines = _collect_between(
            blocks, start_idx, next_end,
            lists_only=bool(list_only_for.get(name, False)),
            include_tables=include_tables_as_lines
        )
        found_sections.append({
            "title": name,
            "level": None if owner_level is None else (owner_level + 1),  # best-guess
            "start_index": start_idx,
            "end_index": next_end,
            "lines": lines
        })

    return {
        "owner_title": owner_title,
        "owner_prefix": owner_pref,
        "owner_family": owner_family,
        "owner_level": owner_level,      # may be None if text-only owner
        "module": module_val,
        "sections": found_sections,
        "bounds": [owner_start, owner_end],
    }


# ------------------------ Hit selection (family preference) -----------

def _infer_preferred_families_from_query(q: str) -> List[str]:
    ql = q.lower()
    prefs: List[str] = []
    if any(w in ql for w in ["container", "cntr", "equipment"]): prefs.append("CNTR")
    if any(w in ql for w in ["vessel", "voyage", "etb", "eta", "baplie", "coprar"]): prefs.append("VSL")
    if any(w in ql for w in ["api", "webhook", "oauth", "token"]): prefs.append("API")
    if any(w in ql for w in ["edi", "edifact", "ansi x12", "iftsta", "coarri", "315", "301", "214"]): prefs.append("EDI")
    seen, out = set(), []
    for p in prefs:
        if p not in seen:
            out.append(p); seen.add(p)
    return out

def pick_best_hit_with_owner(
    results: List[Tuple[str, str, float]],
    docx_path: str,
    preferred_families: Optional[List[str]] = None
) -> Optional[Tuple[str, str, float, Dict[str, Any]]]:
    if not results:
        return None

    # derive each hit's owner prefix quickly via paragraph text
    doc = Document(docx_path)
    blocks: List[Tuple[str, Any]] = list(iter_block_items(doc))
    block_index = build_doc_index(doc)

    def quick_owner_family(sid: str) -> Optional[str]:
        root = _parse_source_id_root(sid)
        if not root or root not in block_index: return None
        bidx = block_index[root]
        kind, obj = blocks[bidx]
        if kind != 'p': return None
        text = (obj.text or "").strip()
        pref = _owner_prefix(text)
        if not pref: return None
        up = pref.upper()
        if up in {"VSL", "VSSL", "VESSEL"}: return "VSL"
        return up

    enriched = []
    for sid, text, score in results:
        fam = quick_owner_family(sid)
        enriched.append((sid, text, score, fam))

    if preferred_families:
        def rank(f): return preferred_families.index(f) if f in preferred_families else 999
        enriched.sort(key=lambda x: (rank(x[3]), -x[2]))
    else:
        enriched.sort(key=lambda x: -x[2])

    # compute full owner bundle for chosen
    sid, text, score, fam = enriched[0]
    owner_bundle = extract_owner_with_sections(docx_path, sid)
    return (sid, text, score, {
        "owner_title": owner_bundle["owner_title"],
        "owner_family": owner_bundle["owner_family"],
        "owner_prefix": owner_bundle["owner_prefix"],
        "bounds": owner_bundle["bounds"]
    })


# ------------------------ JSON assembly -------------------------------

def results_to_json(
    query: str,
    docx_path: str,
    results: List[Tuple[str, str, float]],
    chosen_enriched: Optional[Tuple[str, str, float, Dict[str, Any]]] = None,
    owner_bundle: Optional[Dict[str, Any]] = None,
) -> str:
    payload = {
        "query": query,
        "docx_path": os.path.abspath(docx_path),
        "count": len(results),
        "results": [
            {"source_id": sid, "text": text, "score": float(score)}
            for sid, text, score in results
        ],
        "chosen_match": (
            {
                "source_id": chosen_enriched[0],
                "snippet": (chosen_enriched[1] or "")[:250],
                "score": float(chosen_enriched[2]),
                "owner_detected": chosen_enriched[3],
            } if chosen_enriched else None
        ),
        "owner": owner_bundle or None
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


# ------------------------ Orchestration --------------------------------

def main(
    query: str,
    doc_path: str,
    preferred_owner_families: Optional[List[str]] = None
) -> str:
    
    results = semantic_docx_search(
        query=query,
        docx_path=doc_path,
        threshold=DEFAULT_THRESHOLD,
        top_k=5,
        include_tables=True,
        chunk=True,
        max_chars=500,
        return_always_topk=True,
        keyword_fallback=True
    )

    if preferred_owner_families is None:
        preferred_owner_families = _infer_preferred_families_from_query(query)

    chosen = pick_best_hit_with_owner(results, doc_path, preferred_families=preferred_owner_families)

    owner_bundle = None
    if chosen:
        sid = chosen[0]
        owner_bundle = extract_owner_with_sections(
            docx_path=doc_path,
            hit_source_id=sid,
            wanted_sections=SUBSECTION_TITLES,
            list_only_for={"Overview": False, "Resolution": True, "Verification": True},
            include_tables_as_lines=True
        )

    return results_to_json(
        query, doc_path, results,
        chosen_enriched=chosen,
        owner_bundle=owner_bundle
    )


# -------------------------- CLI --------------------------------------

if __name__ == "__main__":
    query = "EDI: Spike in DLQ messages after routine maintenance; consumer group lag increased across EDI topic"
    DOCX_PATH = os.path.join(
        os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))),
        "Info",
        "Knowledge Base.docx"
    )
    json_str = main(query, DOCX_PATH, preferred_owner_families=None)
    print(json_str)
