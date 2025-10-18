import os, re, json
from datetime import datetime
from difflib import SequenceMatcher
from typing import Any

from LLM_folder.call_openai_basic import ask_gpt5


DEFAULT_EXCEL_PATH = os.path.join(
    os.path.dirname(os.path.dirname(__file__)), "Info", "Case Log.xlsx"
)
DEFAULT_KB_PATH = os.path.join(
    os.path.dirname(os.path.dirname(__file__)), "Info", "Knowledge Base.docx"
)


def _to_json(text: str):
    """Best-effort parse of JSON from model text output."""
    try:
        return json.loads(text)
    except Exception:
        m = re.search(r"\{[\s\S]*\}", text or "")
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                return None
    return None


def extract_email_entities(email_text: str) -> dict:
    """Use LLM to extract structured fields from an email."""
    sys = (
        "Return ONLY compact JSON for fields: "
        "subject, customer, product, environment, "
        "error_codes[], services[], keywords[], case_ids[], "
        "probable_issue, suspected_component, summary."
    )
    user = (
        "Extract fields from email below. If missing, use null or empty arrays.\n" \
        f"Email:\n{email_text}"
    )
    raw = ask_gpt5(user, system_prompt=sys, max_completion_tokens=2048)
    data = _to_json(raw) or {}
    if not data:
        raw = ask_gpt5("ONLY JSON. " + user, system_prompt=sys, max_completion_tokens=2048)
        data = _to_json(raw) or {}
    for k in [
        "subject", "customer", "product", "environment",
        "probable_issue", "suspected_component", "summary",
    ]:
        data.setdefault(k, None)
    for k in ["error_codes", "services", "keywords", "case_ids"]:
        data.setdefault(k, [])
    return data


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip()).lower()


def _toks(s: str):
    s = _norm(s)
    s = re.sub(r"[^a-z0-9_\- ]+", " ", s)
    return [t for t in s.split() if t]


def _ratio(a: str, b: str) -> float:
    return SequenceMatcher(None, _norm(a), _norm(b)).ratio()


def _load_excel(excel_path: str):
    try:
        import pandas as pd
    except Exception as e:
        raise RuntimeError(f"pandas required to read Excel: {e}")
    if not os.path.exists(excel_path):
        raise FileNotFoundError(excel_path)
    xls = pd.ExcelFile(excel_path)
    sheet = xls.sheet_names[0]
    return xls.parse(sheet), sheet


def _canon_cols(cols):
    m = {}
    alt = {
        "case_id": {"case id", "caseid", "id", "ticket", "ticket id"},
        "subject": {"subject", "title", "summary"},
        "description": {"description", "details", "notes", "body"},
        "customer": {"customer", "client", "account"},
        "product": {"product", "service", "module"},
        "error_code": {"error", "error code", "code"},
        "date": {"date", "created", "opened", "created at"},
    }
    def simp(c):
        c = c.strip().lower().replace("_", " ")
        return re.sub(r"\s+", " ", c)
    for c in cols:
        sc = simp(c)
        hit = False
        for k, alts in alt.items():
            if sc in alts or sc == k:
                m[k] = c; hit = True; break
        if not hit:
            m.setdefault(sc, c)
    return m


def find_matching_cases(entities: dict, excel_path: str = DEFAULT_EXCEL_PATH, max_results: int = 10) -> list:
    """Score and return likely matching past cases from the Excel log."""
    df, sheet = _load_excel(excel_path)
    mp = _canon_cols(list(df.columns))
    subj = entities.get("subject") or ""
    prod = entities.get("product") or ""
    kws = entities.get("keywords") or []
    ecs = entities.get("error_codes") or []
    ids = set(_norm(x) for x in (entities.get("case_ids") or []))
    kwt = set(_toks(" ".join([subj, *kws])))
    out = []
    for _, r in df.iterrows():
        d = r.to_dict()
        rs = str(d.get(mp.get("subject", "subject"), ""))
        rd = str(d.get(mp.get("description", "description"), ""))
        rp = str(d.get(mp.get("product", "product"), ""))
        text = f"{rs}\n{rd}"
        sc = 0.0
        cid_col = mp.get("case_id")
        if cid_col and _norm(str(d.get(cid_col, ""))) in ids:
            sc += 100.0
        if prod:
            sc += 8.0 if _norm(prod) == _norm(rp) else 5.0 * _ratio(prod, rp)
        for ec in ecs:
            if ec and re.search(re.escape(str(ec)), text, re.I):
                sc += 6.0
        ov = len(kwt & set(_toks(text)))
        sc += min(ov, 10) * 1.0
        sc += 4.0 * _ratio(subj, rs)
        if sc > 0.5:
            d["_match_score"] = round(sc, 3)
            d["_sheet"] = sheet
            out.append(d)
    out.sort(key=lambda x: x["_match_score"], reverse=True)
    return out[:max_results]


def process_email(email_info: Any, excel_path: str = DEFAULT_EXCEL_PATH, max_results: int = 10) -> dict:
    """High-level pipeline to parse email with LLM, match past cases, and generate Problem Statement + SOP."""
    if isinstance(email_info, dict):
        s = email_info.get("subject") or ""
        b = email_info.get("body") or email_info.get("text") or ""
        email_text = (s + "\n\n" + b).strip() or json.dumps(email_info)
    else:
        email_text = str(email_info)
    ents = extract_email_entities(email_text)
    matches = find_matching_cases(ents, excel_path=excel_path, max_results=max_results)
    kb_used = False
    if not matches:
        print("No historical case log found. Running Knowledge Base docx...")
        kb_obj = _kb_fallback(ents, matches, DEFAULT_KB_PATH)
        if kb_obj:
            llm_analysis = kb_obj
            kb_used = True
        else:
            llm_analysis = generate_problem_and_sop(ents, matches)
    else:
        llm_analysis = generate_problem_and_sop(ents, matches)
    return {
        "structured_email": ents,
        "matches": matches,
        "llm_analysis": llm_analysis,
        "excel_path": excel_path,
        "kb_fallback_used": kb_used,
        "timestamp": datetime.utcnow().isoformat() + "Z",
    }

def _safe_row_summary(row: dict, max_len: int = 500) -> str:
    # Prefer common fields first
    preferred = [
        "case_id", "Case ID", "ticket", "Ticket ID",
        "subject", "Subject", "title", "summary",
        "product", "Product", "service", "Service",
        "error_code", "Error", "code",
        "date", "created", "Created At",
        "description", "details", "notes", "body",
    ]
    items = []
    used = set()
    for k in preferred:
        if k in row and k not in used and not str(k).startswith("_"):
            v = str(row.get(k, ""))
            if v:
                items.append(f"{k}: {v}")
                used.add(k)
    # Fill with remaining keys if space
    if len("; ".join(items)) < max_len:
        for k, v in row.items():
            if k in used or str(k).startswith("_"):
                continue
            sv = str(v)
            if not sv:
                continue
            items.append(f"{k}: {sv}")
            if len("; ".join(items)) >= max_len:
                break
    text = "; ".join(items)
    return text[:max_len]


def generate_problem_and_sop(structured_email: dict, matches: list) -> dict:
    """
    Use the LLM to synthesize a concise Problem Statement and SOP steps
    based on the parsed email and the top matching historical cases.
    Returns a dict with keys: problem_statement, likely_cause, priority, sop, related_case_ids, assumptions, suggested_escalation.
    """
    # Build compact context from matches
    related_ids = []
    context_lines = []
    for m in matches[:5]:
        rid = None
        for key in ("case_id", "Case ID", "id", "ticket", "Ticket ID"):
            if key in m and m.get(key):
                rid = str(m.get(key))
                break
        if rid:
            related_ids.append(rid)
        summary = _safe_row_summary(m)
        context_lines.append(f"- score={m.get('_match_score')} | {summary}")
    context = "\n".join(context_lines) or "- No prior cases matched with high confidence."

    system = (
        "You are a Level 2 Product Ops assistant. "
        "Given an incoming alert and relevant past cases, produce a concise Problem Statement and a clear, actionable SOP for triage/resolution. "
        "Be specific and grounded in the provided context. If data is insufficient, state targeted next-checks. Respond ONLY with minified JSON."
    )
    schema_hint = {
        "problem_statement": "string",
        "likely_cause": "string|null",
        "priority": "P1|P2|P3|Unknown",
        "sop": ["step 1", "step 2", "..."],
        "related_case_ids": ["string"],
        "assumptions": ["string"],
        "suggested_escalation": "string|null",
    }
    user = (
        "Incoming alert (structured):\n" + json.dumps(structured_email, ensure_ascii=False) + "\n\n" +
        "Top matching past cases (compact):\n" + context + "\n\n" +
        "Produce JSON with fields: " + json.dumps(schema_hint) + ". "
        "Keep it brief but complete."
    )

    raw = ask_gpt5(user, system_prompt=system, max_completion_tokens=2048)
    obj = _to_json(raw) or {}
    # Ensure keys
    obj.setdefault("problem_statement", "")
    obj.setdefault("likely_cause", None)
    obj.setdefault("priority", "Unknown")
    obj.setdefault("sop", [])
    obj.setdefault("related_case_ids", related_ids)
    obj.setdefault("assumptions", [])
    obj.setdefault("suggested_escalation", None)
    obj.setdefault("source", "direct")

    # Fallback: If no meaningful output, try leveraging the Knowledge Base
    needs_kb = (not obj.get("problem_statement")) or (not obj.get("sop"))
    if needs_kb:
        kb_obj = _kb_fallback(structured_email, matches, DEFAULT_KB_PATH)
        if kb_obj:
            # Merge keeping non-empty fields
            for k in [
                "problem_statement",
                "likely_cause",
                "priority",
                "sop",
                "related_case_ids",
                "assumptions",
                "suggested_escalation",
            ]:
                if not obj.get(k) and kb_obj.get(k):
                    obj[k] = kb_obj[k]
            obj["source"] = "kb_fallback"
            obj["kb_path"] = DEFAULT_KB_PATH
    return obj


def _read_kb_docx(kb_path: str) -> str:
    try:
        from docx import Document  # python-docx
    except Exception:
        return ""
    if not os.path.exists(kb_path):
        return ""
    try:
        doc = Document(kb_path)
        parts = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        return "\n".join(parts)
    except Exception:
        return ""


def _kb_extract_relevant(kb_text: str, terms: list, max_chars: int = 15000) -> str:
    if not kb_text:
        return ""
    if not terms:
        return kb_text[:max_chars]
    chunks = re.split(r"\n\s*\n+", kb_text)
    sel = []
    terms_l = [t.lower() for t in terms if t]
    for ch in chunks:
        cl = ch.lower()
        if any(t in cl for t in terms_l):
            sel.append(ch)
        if sum(len(s) for s in sel) >= max_chars:
            break
    if not sel:
        return kb_text[:max_chars]
    out = "\n\n".join(sel)
    return out[:max_chars]


def _read_kb_sections(kb_path: str) -> list:
    """
    Read Knowledge Base docx and split into sections using Heading paragraphs.
    Returns a list of dicts: {title, overview, content} where overview is extracted
    from paragraphs labeled 'Overview' or the first paragraph when no explicit
    label exists.
    """
    try:
        from docx import Document
    except Exception:
        return []
    if not os.path.exists(kb_path):
        return []
    try:
        doc = Document(kb_path)
        sections = []
        current = {"title": None, "paras": []}
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if not text:
                continue
            style = getattr(p.style, "name", "") or ""
            if style.lower().startswith("heading"):
                if current["title"] or current["paras"]:
                    sections.append(current)
                current = {"title": text, "paras": []}
            else:
                current["paras"].append(text)
        if current["title"] or current["paras"]:
            sections.append(current)

        # Build overview/content strings
        out = []
        for s in sections:
            title = s.get("title") or "Untitled"
            paras = s.get("paras") or []
            overview = []
            in_overview = False
            for t in paras:
                low = t.lower()
                if low.startswith("overview:") or low == "overview":
                    in_overview = True
                    # Remove the leading label
                    ov = t.split(":", 1)[-1].strip()
                    if ov:
                        overview.append(ov)
                    continue
                # Stop overview if we hit a likely subheading
                if in_overview and (low.endswith(":") or low in {"steps", "procedure", "runbook", "resolution"}):
                    break
                if in_overview:
                    overview.append(t)
            # If no explicit overview, take the first paragraph as a brief overview
            if not overview and paras:
                overview = [paras[0]]
            out.append({
                "title": title,
                "overview": " ".join(overview).strip(),
                "content": "\n".join(paras).strip(),
            })
        return out
    except Exception:
        return []


def _kb_overview_corpus(kb_path: str, terms: list, max_chars: int = 12000) -> str:
    """Assemble a corpus consisting of Title + Overview for matching sections."""
    sections = _read_kb_sections(kb_path)
    if not sections:
        # fallback to plain text mode
        return _kb_extract_relevant(_read_kb_docx(kb_path), terms, max_chars)
    terms_l = [t.lower() for t in (terms or []) if t]
    picks = []
    for s in sections:
        blob = (s.get("title", "") + "\n" + s.get("overview", "")).lower()
        if not terms_l or any(t in blob for t in terms_l):
            title = s.get("title") or "Untitled"
            overview = s.get("overview") or ""
            picks.append(f"Title: {title}\nOverview: {overview}".strip())
        if sum(len(x) for x in picks) >= max_chars:
            break
    if not picks:
        # If nothing matched on titles/overviews, include first few overviews
        for s in sections[:5]:
            title = s.get("title") or "Untitled"
            overview = s.get("overview") or ""
            picks.append(f"Title: {title}\nOverview: {overview}".strip())
            if sum(len(x) for x in picks) >= max_chars:
                break
    corpus = "\n\n".join(picks)
    return corpus[:max_chars]


def _kb_fallback(structured_email: dict, matches: list, kb_path: str) -> dict:
    # Build terms from entities
    terms = []
    terms += structured_email.get("keywords", []) or []
    terms += structured_email.get("services", []) or []
    terms += structured_email.get("error_codes", []) or []
    terms += structured_email.get("case_ids", []) or []
    corpus = _kb_overview_corpus(kb_path, terms)
    if not corpus:
        return {}

    # Use KB content to synthesize Problem + SOP
    system = (
        "You are a Level 2 Product Ops assistant. Using the provided knowledge "
        "base OVERVIEW excerpts and the incoming alert, produce a concise Problem Statement "
        "and an actionable SOP. If KB guidance is not directly applicable, infer "
        "best practices clearly. Respond ONLY with minified JSON."
    )
    schema_hint = {
        "problem_statement": "string",
        "likely_cause": "string|null",
        "priority": "P1|P2|P3|Unknown",
        "sop": ["step 1", "step 2", "..."],
        "related_case_ids": ["string"],
        "assumptions": ["string"],
        "suggested_escalation": "string|null",
    }
    user = (
        "Incoming alert (structured):\n" + json.dumps(structured_email, ensure_ascii=False) + "\n\n" +
        "Relevant KB Overviews (Title + Overview only):\n" + corpus + "\n\n" +
        "Produce JSON with fields: " + json.dumps(schema_hint) + ". "
    )
    raw = ask_gpt5(user, system_prompt=system, max_completion_tokens=2048)
    kb_obj = _to_json(raw) or {}
    if not kb_obj:
        return {}
    kb_obj.setdefault("problem_statement", "")
    kb_obj.setdefault("likely_cause", None)
    kb_obj.setdefault("priority", "Unknown")
    kb_obj.setdefault("sop", [])
    kb_obj.setdefault("related_case_ids", [])
    kb_obj.setdefault("assumptions", [])
    kb_obj.setdefault("suggested_escalation", None)
    kb_obj["source"] = "kb_fallback"
    return kb_obj


# === New JSON alert entrypoint ===
def _extract_codes(text: str) -> list:
    """Find code-like tokens (e.g., ALR-123, INC-123456, REF-ABC-0001, CORR-0001)."""
    if not text:
        return []
    pat = r"\b(?:ALR|INC|TCK|REF|CORR|ERR|ORA|SQL|HTTP|VR|VRR|COPARN|COARRI|BAPLIE)[-_]?[A-Z0-9]+(?:[-_][A-Z0-9]+)*\b"
    return sorted(set(re.findall(pat, text, flags=re.IGNORECASE)), key=str.lower)


def entities_from_alert(alert: dict) -> dict:
    """
    Map a structured alert JSON to the matcher 'entities' shape.
    Supports old and new keys: container_ids|cntr_no, edi_refs|message_ref,
    vessel_names|vessel_name, correlation_id.
    """
    variables = alert.get("variables", {}) or {}
    evidence = alert.get("evidence", {}) or {}
    text_sample = (evidence.get("text_sample") or "") if isinstance(evidence, dict) else ""

    def get_list(*keys):
        out = []
        for k in keys:
            v = variables.get(k)
            if isinstance(v, list):
                out.extend([str(x) for x in v if x])
            elif isinstance(v, str) and v:
                out.append(v)
        return out

    # Build keywords from variables and flags (supports new names)
    kw = []
    kw += get_list("container_ids", "cntr_no")
    kw += get_list("edi_refs", "message_ref")
    kw += get_list("vessel_names", "vessel_name")
    kw += get_list("voyages")
    kw += get_list("terminals")
    kw += get_list("edi_types")
    kw += get_list("correlation_id")
    if variables.get("is_duplicate_hint"):
        kw.append("duplicate")
    if variables.get("is_ack_missing_hint"):
        kw.extend(["ack missing", "missing ack", "no ack"])

    # From evidence text
    kw += _extract_codes(text_sample)

    # Services: use edi_types and incident type
    services = []
    for s in get_list("edi_types"):
        services.append(s)
    if alert.get("incident_type"):
        services.append(str(alert.get("incident_type")))

    # Error codes
    error_codes = get_list("error_codes")

    # Case ids: include message refs, correlation ids, and extracted codes
    case_ids = []
    case_ids += get_list("edi_refs", "message_ref")
    case_ids += get_list("correlation_id")
    case_ids += _extract_codes(text_sample)

    subject = alert.get("problem_statement") or (alert.get("incident_type") or "Alert")
    summary = subject

    entities = {
        "subject": subject,
        "customer": None,
        "product": alert.get("incident_type") or None,
        "environment": None,
        "error_codes": error_codes,
        "services": services,
        "keywords": kw,
        "case_ids": case_ids,
        "probable_issue": ("ACK missing" if variables.get("is_ack_missing_hint") else None),
        "suspected_component": None,
        "summary": summary,
        "source_alert": alert,
    }
    return entities


def process_alert_json(alert: Any, excel_path: str = DEFAULT_EXCEL_PATH, max_results: int = 10) -> dict:
    """
    Accepts an already-structured alert JSON like the example provided.
    Builds matching entities, finds matching Excel cases, and uses the LLM
    to generate a Problem Statement and SOP.
    """
    if isinstance(alert, str):
        try:
            alert_obj = json.loads(alert)
        except Exception:
            alert_obj = {"problem_statement": alert}
    else:
        alert_obj = alert or {}

    entities = entities_from_alert(alert_obj)
    matches = find_matching_cases(entities, excel_path=excel_path, max_results=max_results)
    kb_used = False
    if not matches:
        print("No historical case log found. Running Knowledge Base docx...")
        kb_obj = _kb_fallback(entities, matches, DEFAULT_KB_PATH)
        if kb_obj:
            llm_analysis = kb_obj
            kb_used = True
        else:
            llm_analysis = generate_problem_and_sop(entities, matches)
    else:
        llm_analysis = generate_problem_and_sop(entities, matches)
    return {
        "structured_alert": entities,
        "matches": matches,
        "llm_analysis": llm_analysis,
        "excel_path": excel_path,
        "kb_fallback_used": kb_used,
        "timestamp": datetime.utcnow().isoformat() + "Z",
    }
